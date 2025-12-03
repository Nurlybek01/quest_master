import os
import qrcode
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Inches


TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
PARCHMENTS_DIR = Path("./parchments")
PARCHMENTS_DIR.mkdir(exist_ok=True)

# Jinja2 окружение
env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))


def render_template(template_name: str, context: Dict[str, Any]) -> str:
    template = env.get_template(template_name)
    return template.render(**context)


def export_to_pdf(quest_data: Dict[str, Any], template: str = "royal_decree.html", with_qr: bool = False) -> Path:
    """Экспортирует квест в PDF с опциональным QR-кодом."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{quest_data['id']}_{timestamp}.pdf"
    output_path = PARCHMENTS_DIR / filename

    context = {
        "quest": quest_data,
        "current_date": datetime.now().strftime("%d.%m.%Y")
    }

    html_content = render_template(template, context)

    if with_qr:
        # Генерируем URL квеста
        url = f"http://quest.local/view/{quest_data['id']}"
        qr = qrcode.make(url)
        qr_filename = f"qr_{quest_data['id']}_{timestamp}.png"
        qr_path = PARCHMENTS_DIR / qr_filename
        qr.save(qr_path)

        # Вставляем QR в HTML (относительный путь — WeasyPrint работает в ./)
        qr_tag = f'<div style="text-align:center; margin-top:20px;"><img src="{qr_filename}" width="100" alt="QR-код"></div>'
        html_content = html_content.replace("</body>", qr_tag + "\n</body>")

    # Генерация PDF
    HTML(string=html_content, base_url=str(PARCHMENTS_DIR)).write_pdf(output_path)

    # Удаляем временный QR только если он не нужен после (опционально, но чисто)
    if with_qr and qr_path.exists():
        qr_path.unlink()

    return output_path


def export_to_docx(quest_data: Dict[str, Any], template: str = "guild_contract.html", with_qr: bool = False) -> Path:
    """Экспортирует квест в DOCX с опциональным QR-кодом."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{quest_data['id']}_{timestamp}.docx"
    output_path = PARCHMENTS_DIR / filename

    doc = Document()
    doc.add_heading(f"Квест #{quest_data['id']}", 0)
    doc.add_paragraph(f"Название: {quest_data['title']}")
    doc.add_paragraph(f"Сложность: {quest_data['difficulty']}")
    doc.add_paragraph(f"Награда: {quest_data['reward']} золотых")
    doc.add_paragraph(f"Описание: {quest_data['description']}")
    doc.add_paragraph(f"Срок выполнения: {quest_data['deadline']}")
    doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}")

    if with_qr:
        url = f"http://quest.local/view/{quest_data['id']}"
        qr = qrcode.make(url)
        qr_path = PARCHMENTS_DIR / f"qr_docx_{quest_data['id']}_{timestamp}.png"
        qr.save(qr_path)
        doc.add_paragraph("QR-код квеста:")
        doc.add_picture(str(qr_path), width=Inches(1.5))
        # Удалять QR для DOCX не обязательно — он встроен в файл

    doc.save(output_path)
    return output_path


class TemplateEngine:
    @staticmethod
    def export(quest_data: Dict[str, Any], format: str = "pdf", template: str = "royal_decree.html", with_qr: bool = False):
        if format == "pdf":
            return export_to_pdf(quest_data, template, with_qr)
        elif format == "docx":
            return export_to_docx(quest_data, template, with_qr)
        else:
            raise ValueError("Поддерживаемые форматы: 'pdf', 'docx'")